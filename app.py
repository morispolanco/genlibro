import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor

# Obtener las claves de API de los secretos de Streamlit
together_api_key = st.secrets["TOGETHER_API_KEY"]

# Función para hacer llamadas a la API de Together
def together_complete(prompt, max_tokens=500):
    url = "https://api.together.xyz/v1/completions"
    headers = {
        "Authorization": f"Bearer {together_api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "meta-llama/Meta-Llama-3.1-405B-Instruct-Turbo",
        "prompt": prompt,
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1.1,  # Penalización para evitar repeticiones
        "stop": ["<|eot_id|>"]
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        return response.json()['choices'][0]['text'].strip()
    else:
        return f"Error: {response.status_code}, {response.text}"

# Función para generar el contenido de cada capítulo con al menos siete páginas simuladas
def generate_chapter_content(chapter_title, min_pages=7):
    page_word_count = 350  # Aproximación de palabras por página
    min_word_count = page_word_count * min_pages  # Palabras mínimas por capítulo
    
    prompt_content = f"Escribe un contenido extenso para el capítulo titulado '{chapter_title}' en el libro. El contenido debe tener al menos {min_word_count} palabras, evitar repeticiones y ser coherente, sin divisiones internas."
    
    # Dividir la generación de contenido en varias llamadas si es necesario
    content = ""
    while len(content.split()) < min_word_count:
        new_content = together_complete(prompt_content, max_tokens=1000)  # Generar bloques grandes de texto
        content += " " + new_content
    
    return content.strip()

# Función para generar el documento DOCX basado en la estructura y contenido con formateo
def generate_formatted_docx(title, thesis_structure, thesis_content):
    doc = Document()
    
    # Formato del título del libro
    title_heading = doc.add_heading(title, 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilos personalizados para los capítulos y párrafos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Aplicar formato a los capítulos y el contenido
    for section in thesis_structure:
        # Formato del título de cada capítulo
        chapter_title = doc.add_heading(section['title'], level=1)
        chapter_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        chapter_title.style.font.name = 'Arial'
        chapter_title.style.font.size = Pt(14)
        chapter_title.style.font.bold = True
        chapter_title.style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Formato del contenido de cada capítulo
        paragraph = doc.add_paragraph(thesis_content[section['title']])
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(18)  # Espaciado entre líneas
        paragraph_format.space_after = Pt(12)   # Espacio después del párrafo
    
    # Guardar el documento en un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Función para generar contenido simultáneamente utilizando multithreading
def generate_all_chapters_concurrently(thesis_structure):
    thesis_content = {}
    
    # Usamos un ThreadPoolExecutor para ejecutar en paralelo
    with ThreadPoolExecutor() as executor:
        # Lanzar tareas simultáneamente
        future_to_section = {executor.submit(generate_chapter_content, section['title']): section for section in thesis_structure}
        
        # Procesar los resultados a medida que se completan
        for future in future_to_section:
            section = future_to_section[future]
            try:
                thesis_content[section['title']] = future.result()
            except Exception as exc:
                thesis_content[section['title']] = f"Error al generar contenido: {exc}"
    
    return thesis_content

st.title("Generación Simultánea de Libros con Exportación Formateada a DOCX")

# Sección de Generación de Libro
st.header("Generación de Libro")
book_title = st.text_input("Ingresa el título de tu libro:")
book_topic = st.text_input("Ingresa el tema central de tu libro (puede ser cualquier tema):")

if st.button("Generar Libro y Estructura"):
    # Generar la introducción y estructura del libro
    prompt_thesis = f"Proporciona una propuesta de libro sobre el tema: '{book_topic}'."
    book_statement = together_complete(prompt_thesis)
    st.subheader("Propuesta del Libro")
    st.write(book_statement)
    
    # Generar la estructura automáticamente basada en el tema
    prompt_structure = f"Con base en el siguiente libro: '{book_statement}', propone una tabla de contenidos detallada para el libro, sin divisiones internas en los capítulos. Incluye una introducción y un número adecuado de capítulos que cubran el tema de manera exhaustiva."
    structure_response = together_complete(prompt_structure)
    st.subheader("Estructura Propuesta")
    st.write(structure_response)
    
    # Procesar la estructura en un formato que se pueda utilizar
    thesis_structure = [{"title": chapter.strip()} for chapter in structure_response.split("\n") if chapter.strip()]
    
    # Generar contenido extenso para cada capítulo de forma simultánea
    st.subheader("Generando contenido extenso para cada capítulo de forma simultánea...")
    thesis_content = generate_all_chapters_concurrently(thesis_structure)
    
    # Exportar a DOCX con formateo
    docx_buffer = generate_formatted_docx(book_title, thesis_structure, thesis_content)
    
    # Descargar el archivo DOCX
    st.download_button(
        label="Descargar Libro en DOCX Formateado",
        data=docx_buffer,
        file_name=f"{book_title}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
